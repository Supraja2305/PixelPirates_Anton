# parsers/word_parser.py
# ============================================================
# Microsoft Word (.docx, .doc) document parser
# Extracts text, tables, and metadata from Office documents
# ============================================================

import io
import logging
import hashlib
from typing import Optional, Dict, List, Any
from pathlib import Path

from ..utils.error_handler import log_error, PipelineError


logger = logging.getLogger(__name__)


class WordParser:
    """
    Parse Microsoft Word documents (.docx, .doc).
    Extracts text, tables, formatting, and metadata.
    """
    
    # Supported extensions
    SUPPORTED_FORMATS = {'.docx', '.doc', '.odt'}
    
    # Magic bytes for Word documents
    MAGIC_BYTES = {
        b'PK\x03\x04': 'docx',  # ZIP-based (Office Open XML)
        b'\xd0\xcf\x11\xe0': 'doc',   # OLE2 (legacy Word)
    }

    def __init__(self):
        """Initialize Word parser."""
        self.logger = logging.getLogger(self.__class__.__name__)

    def extract_text_from_word(self, file_bytes: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Extract text and metadata from Word document.
        
        Args:
            file_bytes: Raw document bytes
            filename: Optional filename for validation
        
        Returns:
            Dict with keys:
                - text: Extracted full text
                - tables: List of extracted tables
                - pages: Page count estimate
                - metadata: Document metadata
                - confidence: Extraction confidence (0.0-1.0)
                - method: 'python-docx' or 'docx2txt'
                - char_count: Total characters
                - checksum: SHA256 of file
                - errors: List of extraction errors
        """
        try:
            # Validate file
            if not file_bytes or len(file_bytes) == 0:
                raise PipelineError("Empty Word document provided")
            
            # Detect format
            doc_format = self._detect_format(file_bytes)
            if not doc_format:
                raise PipelineError(f"Unrecognized Word document format for {filename}")
            
            # Try python-docx for .docx files
            if doc_format == 'docx':
                return self._extract_docx(file_bytes, filename)
            
            # For .doc files, try python-docx or fallback
            elif doc_format == 'doc':
                return self._extract_doc(file_bytes, filename)
            
            # For other formats
            else:
                raise PipelineError(f"Unsupported document format: {doc_format}")
                
        except Exception as e:
            error_msg = f"Word extraction failed: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            log_error(error_msg, context={"filename": filename})
            raise PipelineError(error_msg)

    def _detect_format(self, file_bytes: bytes) -> Optional[str]:
        """Detect Word document format by magic bytes."""
        for magic, fmt in self.MAGIC_BYTES.items():
            if file_bytes.startswith(magic):
                return fmt
        return None

    def _extract_docx(self, file_bytes: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Extract from modern Word (.docx) using python-docx.
        DOCX is ZIP-based XML format (Office Open XML).
        """
        try:
            from docx import Document
            from docx.table import Table
        except ImportError:
            raise PipelineError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Load document from bytes
            doc_stream = io.BytesIO(file_bytes)
            doc = Document(doc_stream)
            
            # Extract full text
            text_parts = []
            tables_data = []
            
            for element in doc.element.body:
                # Extract paragraphs
                if element.tag.endswith('p'):
                    para = element.getparent().find('.//w:p', namespaces=element.nsmap or {})
                    if para is not None:
                        text_parts.append(para.text if hasattr(para, 'text') else '')
                
                # Extract tables
                if element.tag.endswith('tbl'):
                    table_data = self._extract_table_from_docx(element)
                    if table_data:
                        tables_data.append(table_data)
            
            # Simpler approach: iterate through paragraphs and tables
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = self._extract_table_from_docx(table)
                if table_data:
                    tables_data.append(table_data)
            
            # Combine text
            full_text = '\n'.join(text_parts)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or None,
                "author": core_props.author or None,
                "subject": core_props.subject or None,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "pages": len(doc.paragraphs) // 50 + 1,  # Rough estimate
            }
            
            # Calculate metrics
            char_count = len(full_text)
            confidence = 0.95 if char_count > 50 else 0.7
            
            checksum = hashlib.sha256(file_bytes).hexdigest()
            
            return {
                "text": full_text,
                "tables": tables_data,
                "pages": metadata.get("pages", 1),
                "metadata": metadata,
                "confidence": confidence,
                "method": "python-docx",
                "char_count": char_count,
                "checksum": checksum,
                "errors": []
            }
            
        except Exception as e:
            self.logger.error(f"DOCX extraction error: {str(e)}")
            raise

    def _extract_doc(self, file_bytes: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Extract from legacy Word (.doc) files.
        Uses python-docx if it supports OLE2, otherwise basic extraction.
        """
        try:
            # Try using docx library as fallback
            from docx import Document
            doc_stream = io.BytesIO(file_bytes)
            
            try:
                doc = Document(doc_stream)
                # If successful, process like DOCX
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                full_text = '\n'.join(text_parts)
                
                tables_data = []
                for table in doc.tables:
                    table_data = self._extract_table_from_docx(table)
                    if table_data:
                        tables_data.append(table_data)
                
                char_count = len(full_text)
                confidence = 0.85 if char_count > 50 else 0.6
                
                checksum = hashlib.sha256(file_bytes).hexdigest()
                
                return {
                    "text": full_text,
                    "tables": tables_data,
                    "pages": len(text_parts) // 40 + 1,
                    "metadata": {"format": "doc", "warning": "Legacy format - limited metadata extraction"},
                    "confidence": confidence,
                    "method": "python-docx-fallback",
                    "char_count": char_count,
                    "checksum": checksum,
                    "errors": ["Legacy .doc format - recommend converting to .docx for better extraction"]
                }
            except Exception:
                # Fallback to basic binary extraction
                return self._extract_doc_basic(file_bytes, filename)
                
        except ImportError:
            return self._extract_doc_basic(file_bytes, filename)

    def _extract_doc_basic(self, file_bytes: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Basic extraction from .doc without library (fallback).
        Extracts readable ASCII/Unicode strings from binary.
        """
        try:
            # Extract text-like sequences
            text_chunks = []
            max_chunk_size = 256
            current_chunk = b''
            
            for byte in file_bytes:
                # Keep printable ASCII and common Unicode patterns
                if 32 <= byte <= 126 or byte in [9, 10, 13]:  # printable + tab, newline, CR
                    current_chunk += bytes([byte])
                else:
                    if len(current_chunk) > 5:
                        try:
                            text = current_chunk.decode('utf-8', errors='ignore')
                            if text.strip():
                                text_chunks.append(text)
                        except:
                            pass
                    current_chunk = b''
            
            full_text = '\n'.join(text_chunks)
            char_count = len(full_text)
            
            checksum = hashlib.sha256(file_bytes).hexdigest()
            
            return {
                "text": full_text,
                "tables": [],
                "pages": char_count // 4000 + 1,
                "metadata": {"format": "doc", "extraction": "binary-fallback"},
                "confidence": 0.5,  # Low confidence for basic extraction
                "method": "binary-extraction",
                "char_count": char_count,
                "checksum": checksum,
                "errors": ["Using basic binary extraction - install python-docx for better quality"]
            }
        except Exception as e:
            self.logger.error(f"DOC binary extraction failed: {str(e)}")
            raise PipelineError(f"Could not extract from .doc file: {str(e)}")

    def _extract_table_from_docx(self, table) -> Optional[Dict[str, Any]]:
        """Extract table structure and content from DOCX table object."""
        try:
            rows = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text (may have multiple paragraphs)
                    cell_text = '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])
                    row_data.append(cell_text)
                rows.append(row_data)
            
            if not rows:
                return None
            
            # Format as markdown-style table
            table_text = self._format_table(rows)
            
            return {
                "headers": rows[0] if rows else [],
                "rows": rows[1:] if len(rows) > 1 else [],
                "markdown": table_text,
                "row_count": len(rows)
            }
        except Exception as e:
            self.logger.warning(f"Table extraction issue: {str(e)}")
            return None

    def _format_table(self, rows: List[List[str]]) -> str:
        """Format table rows as markdown."""
        if not rows:
            return ""
        
        lines = []
        for i, row in enumerate(rows):
            # Escape pipe characters and create markdown row
            cells = [' | '.join(cell.split('\n')) for cell in row]
            lines.append(' | '.join(cells))
            
            # Add header separator after first row
            if i == 0:
                sep = ' | '.join(['---'] * len(row))
                lines.insert(1, sep)
        
        return '\n'.join(lines)

    def validate_word_file(self, file_bytes: bytes, filename: str = None) -> tuple:
        """
        Validate Word document without full extraction.
        
        Returns:
            (is_valid: bool, format: str, error: str or None)
        """
        try:
            if not file_bytes:
                return (False, None, "Empty file")
            
            fmt = self._detect_format(file_bytes)
            if not fmt:
                return (False, None, "Not a recognized Word document")
            
            if fmt not in self.SUPPORTED_FORMATS:
                return (False, fmt, f"Format {fmt} not yet supported")
            
            return (True, fmt, None)
        except Exception as e:
            return (False, None, str(e))
